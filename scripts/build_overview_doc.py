"""Generate docs/PROTOTYPE_OVERVIEW.docx — a complete walkthrough of the
IMM Digital Twin Platform. Screenshots are read from docs/screenshots/platform/.

Run:  .venv/bin/python scripts/build_overview_doc.py
"""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
SHOTS = ROOT / "docs" / "screenshots" / "platform"
OUT = ROOT / "docs" / "PROTOTYPE_OVERVIEW.docx"


def H(doc, text, level=1):
    doc.add_heading(text, level=level)


def P(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    return p


def bullet(doc, text):
    doc.add_paragraph(text, style="List Bullet")


def code(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(9)


def img(doc, name, caption, width=6.5):
    path = SHOTS / name
    if not path.exists():
        P(doc, f"[missing screenshot: {name}]", italic=True)
        return
    doc.add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(f"Figure: {caption}")
    r.italic = True
    r.font.size = Pt(9)


def _shade(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    cell._tc.get_or_add_tcPr().append(shd)


def glossary_entry(doc, term, technical, plain, floor):
    p = doc.add_paragraph()
    p.add_run(f"{term} — ").bold = True
    p.add_run(technical)
    pl = doc.add_paragraph()
    pl.add_run("In plain English: ").bold = True
    pl.add_run(plain)
    fl = doc.add_paragraph()
    r = fl.add_run("On the factory floor: ")
    r.bold = True
    r.italic = True
    fl.add_run(floor).italic = True
    doc.add_paragraph()


def palette_table(doc, rows):
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for c, txt in zip(table.rows[0].cells, ("Token", "Swatch", "Hex · purpose")):
        c.paragraphs[0].add_run(txt).bold = True
    for token, hexv, purpose in rows:
        cells = table.add_row().cells
        cells[0].text = token
        _shade(cells[1], hexv)
        cells[2].text = f"{hexv} · {purpose}"


def subsystem_table(doc):
    rows = [
        ("Hydraulic", "#3b82f6", "hydraulic", "Pumps, valves, cylinders, oil circuit", "Inj. pressure, oil temp"),
        ("Screw & Check Ring", "#f97316", "screw_check_ring", "Screw, non-return ring, barrel, nozzle", "Cushion, cavity pressure"),
        ("Drive", "#22c55e", "drive", "Motors, gearboxes, servo drive", "Screw RPM, cycle time"),
        ("Heaters", "#eab308", "heaters", "Barrel heater bands, thermocouples", "Barrel temps, nozzle temp"),
        ("Mold & Clamp", "#a855f7", "mold", "Platens, cavities, ejector, clamp", "Clamp force, cooling time"),
    ]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    for c, txt in zip(table.rows[0].cells, ("Subsystem", "Colour", "Backend key", "What it is", "Key sensors")):
        c.paragraphs[0].add_run(txt).bold = True
    for label, hexv, key, what, sensors in rows:
        cells = table.add_row().cells
        cells[0].text = label
        _shade(cells[1], hexv)
        cells[1].paragraphs[0].add_run(hexv).font.size = Pt(8)
        cells[2].text = key
        cells[3].text = what
        cells[4].text = sensors


def main():
    doc = Document()

    # ---------- Title ----------
    t = doc.add_heading("IMM Digital Twin Platform", 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run("Industrial Predictive-Maintenance Platform for Injection-Molding Machines\n"
                "Prototype Overview — what it is, how it works, and how to demo it").italic = True

    # ---------- 1. Executive summary ----------
    H(doc, "1. Executive summary", 1)
    P(doc,
      "The IMM Digital Twin Platform is an enterprise-grade monitoring environment for injection-"
      "molding machines. A large, interactive 3-D digital twin sits at the centre of a command-"
      "centre layout; each of the machine's five monitored subsystems is colour-coded and tinted "
      "by its live health directly on the model. The platform predicts Remaining Useful Life (RUL) "
      "per subsystem, converts it into calendar replacement dates and urgency, and surfaces it all "
      "with charts and tooltips that explain themselves — so a non-specialist can read the machine "
      "and know what to do.")
    P(doc, "The design target is professional industrial software comparable to platforms from "
           "Siemens, Bosch, Schneider Electric, Rockwell, GE Digital, ABB and PTC ThingWorx — not a "
           "consumer dashboard.")
    img(doc, "01_operations_overview.png",
        "Command Center — Operations mode. The 3-D twin is the hero; the left rail shows live "
        "per-subsystem health; the right panel is the machine-overview context.")

    P(doc, "Two modes, one application:", bold=True)
    bullet(doc, "Operations — live monitoring: health-tinted twin, subsystem rail, contextual detail "
                "panel, enterprise charts, and self-explaining tooltips.")
    bullet(doc, "Inspection — the full mesh-level engineering workbench folded in as a native module: "
                "hierarchy, isolate, wireframe / x-ray, search, and a component-mapping editor.")

    # ---------- 2. Architecture ----------
    H(doc, "2. Architecture", 1)
    P(doc, "The platform is a single-page application served by one Python (aiohttp) server. "
           "A WebSocket streams a machine snapshot every cycle; a small REST surface handles "
           "commands and persists configuration and the component map to files. Nine modules:")
    for n, name, role in [
        ("1", "Digital Twin Viewer", "R3F canvas, camera (orbit/pan/zoom/fit/focus/reset, cinematic transitions), lighting, outline FX"),
        ("2", "Component Mapping Engine", "classifies meshes → subsystems (spatial + name hints); loads/persists the component map"),
        ("3", "Component Health Engine", "per-subsystem health, status band and trend from the live snapshot"),
        ("4", "RUL Engine", "per-subsystem remaining-useful-life, urgency and forecast band"),
        ("5", "Sensor Visualization Engine", "scalar / curve channels mapped to subsystems, with adaptive threshold envelopes"),
        ("6", "Alert Engine", "machine-state banner, fault list, prioritised 'act on this first'"),
        ("7", "Analytics Engine", "trends, forecast regions, predicted-failure markers, KPIs"),
        ("8", "Deployment Module", "config + threshold + map persistence, deterministic boot"),
        ("9", "Collaboration Module", "shared review sessions via URL state, tunnel-aware"),
    ]:
        bullet(doc, f"{n}. {name} — {role}.")
    P(doc, "The five monitored subsystems and how they bind to the backend and the sensors:")
    subsystem_table(doc)
    P(doc, "Data contract (unchanged from the underlying predictive-maintenance engine): the snapshot "
           "carries cycle index, per-component health [0..1], per-component RUL quantiles "
           "(p10/p50/p90) + replacement date + urgency, a 3-class quality prediction, 14 scalar and "
           "5 curve sensor channels, and the latched machine state.")

    # ---------- 3. Operations mode ----------
    H(doc, "3. Operations mode", 1)

    H(doc, "3.1 The Subsystem Rail", 2)
    P(doc, "The left rail lists the five subsystems, each with its identity colour, a live health bar, "
           "status band (Healthy / Watch / Critical / Failed), trend arrow and replacement horizon. "
           "Hovering a subsystem highlights its meshes on the twin and opens an explanation; clicking "
           "it selects the subsystem, flies the camera in, and makes the right panel that subsystem's "
           "live detail. A 'Machine health' rollup at the bottom names the worst component.")
    img(doc, "02_rail_tooltip.png",
        "Hovering a rail item highlights its region on the 3-D model and opens a What / Reading / "
        "Action explanation.")

    H(doc, "3.2 Subsystem selection — the active context", 2)
    P(doc, "Selecting any subsystem (from the rail or by clicking it on the 3-D model) ghosts the rest "
           "of the machine, frames the camera on the selected region, shows a 'Viewing ·' badge for "
           "orientation, and turns the right panel into that subsystem's detail: a health ring, RUL "
           "(p50 cycles + calendar horizon), failure probability, status, trend, a maintenance "
           "recommendation, an RUL-forecast chart, a health-trend chart, and a live sensor summary.")
    img(doc, "03_select_hydraulic.png", "Hydraulic selected — base/power-pack region highlighted; full live detail panel.")
    img(doc, "05_select_drive.png", "Drive selected — the central drive region; the camera reframes to it.")
    img(doc, "07_select_mold.png", "Mold & Clamp selected — the clamp end; per-subsystem charts update to Mold.")

    H(doc, "3.3 Enterprise charts", 2)
    P(doc, "Every chart carries the same enterprise toolkit: a crosshair with hover values, shaded "
           "warning / critical threshold regions, a forecast region with the p10–p90 uncertainty band "
           "around the p50, and a predicted-failure marker. The RUL forecast projects the subsystem's "
           "health decline to its failure threshold; the crosshair reads out the value, the calendar "
           "horizon and the risk band at any point.")
    img(doc, "09_forecast_crosshair.png",
        "RUL Forecast crosshair — '+1.5k cycles · ~39 days / p10–p90 band / p50 health 37% (watch)'.")

    H(doc, "3.4 The self-explaining platform", 2)
    P(doc, "Intelligent tooltips are a first-class feature. Hovering anything answers three questions: "
           "What am I seeing? · Is it good or bad? · What should I do? The judgement is derived from the "
           "live value and its threshold band, so the words always match the numbers.")
    img(doc, "08_kpi_tooltip.png", "A metric tooltip — defines RUL and interprets the current reading.", width=6.5)
    img(doc, "10_sensor_tooltip.png",
        "A sensor tooltip — value, normal direction, band, and the operational meaning (with a live "
        "trend sparkline and normal envelope).")

    H(doc, "3.5 Demonstration controls & settings", 2)
    P(doc, "Demo controls (inject a fault per subsystem, accelerate the clock) are fenced and clearly "
           "labelled as not real machine controls. Settings expose the production rate (cycles/day) "
           "used to turn predicted cycles into calendar dates; the value persists to disk.")
    img(doc, "11_demo_controls.png", "Demonstration controls — fault buttons per subsystem + speed, clearly fenced.", width=6.0)
    img(doc, "12_settings.png", "Settings — production rate, self-explaining and persisted.", width=5.5)

    # ---------- 4. Degradation & failure ----------
    H(doc, "4. Degradation and failure behaviour", 1)
    P(doc, "Inject a fault (or connect a real machine) and the story unfolds live and on the geometry. "
           "As a subsystem wears, its region on the 3-D model shifts green → amber → red, the rail and "
           "machine state escalate Running → Warning → Critical, the RUL forecast collapses toward the "
           "failure line, predicted quality degrades, and alerts appear. When a component crosses its "
           "failure threshold the machine latches FAILED with a banner naming the part and the cycle.")
    img(doc, "13_state_warning.png", "Warning — the Hydraulic base shifts amber; state = Warning.")
    img(doc, "14_state_critical.png", "Critical — the region reddens further; quality flips toward WASTE.")
    img(doc, "15_state_failed.png", "Failed — the region is red, a banner names the failed part and cycle, alerts raised.")

    # ---------- 5. Component mapping engine ----------
    H(doc, "5. Component Mapping Engine", 1)
    P(doc, "The platform binds live health to geometry through a component map (subsystem → mesh ids) "
           "built from the COLLADA model. Because the model's ~660 meshes are mostly generically named, "
           "the classifier combines two deterministic signals: name hints for the few descriptive nodes "
           "(hopper, platen, oil_cooler, …), and spatial zones for the rest — each mesh centroid is "
           "projected onto the machine's principal axis and vertical and mapped to the region that best "
           "matches a real injection-molding layout (power pack at the base, screw/heaters/drive across "
           "the top, mold at the clamp end). Structure (frames, guards, panels, fasteners) is routed to "
           "a neutral 'Structure' bucket so monitored assets stand out.")
    P(doc, "The result is a complete, reproducible partition committed as the source of truth at "
           "public/map/component-map.detailed.json (every mesh in exactly one subsystem). On boot the "
           "platform loads this file; if it is ever absent it derives the map at load time and persists "
           "it back, so the mapping never needs to be recreated. A domain expert can refine it in "
           "Inspection mode and re-export.")
    code(doc, "counts (this build): Hydraulic 219 · Screw&CheckRing 111 · Drive 57 · Heaters 71 · "
              "Mold 184 · Structure 18   (660 meshes total)")

    # ---------- 6. Inspection mode ----------
    H(doc, "6. Inspection mode", 1)
    P(doc, "The original mapping workbench is folded in as a native module — power users inspect "
           "internals without leaving the platform, sharing one scene and camera with Operations.")
    img(doc, "16_inspection_overview.png",
        "Inspection — hierarchy explorer (left), the model in materials view, inspector + mapping "
        "editor (right) with live subsystem counts.")
    bullet(doc, "Mesh selection — click any mesh (or a hierarchy node); the inspector shows its name, "
                "material, triangle/vertex counts, parent and subsystem assignment.")
    bullet(doc, "Isolation — hide everything except the selection to study a part in context.")
    bullet(doc, "Wireframe / X-ray / Edges — see internal structure and overlapping geometry.")
    bullet(doc, "Search — filter the hierarchy by component or material name.")
    bullet(doc, "Mapping editor — assign meshes to subsystems (keys 1–5), reset, export the detailed "
                "map and persist it back to the server.")
    img(doc, "17_inspection_mesh_selected.png", "A mesh selected — the inspector and hierarchy reflect it; outline highlight.")
    img(doc, "18_inspection_wireframe.png", "Wireframe — internal structure of the machine.")
    img(doc, "20_inspection_isolate.png", "Isolation — only the selected part is shown; the inspector details it.")

    # ---------- 7. Collaboration ----------
    H(doc, "7. Collaborative review", 1)
    P(doc, "The current view — mode and selected subsystem — is encoded in the URL hash, so a link "
           "reproduces a colleague's exact view. Combined with the one-command public tunnel, a "
           "developer and a domain expert can review the same live machine remotely:")
    code(doc, "https://<your-tunnel>.trycloudflare.com/#mode=operations&sub=Mold")
    P(doc, "Everyone streams the same live snapshots from the one backend; no extra infrastructure is "
           "required for a shared review session.")
    P(doc, "For investor / customer demos, append &ro=1 for a read-only link: live monitoring stays "
           "fully interactive, but machine controls (reset, settings, fault injection) are hidden so a "
           "remote viewer can explore without disrupting the live machine. The presenter uses the plain "
           "URL for full control. See docs/DGX_DEMO_CHECKLIST.md.")
    img(doc, "23_readonly_review.png",
        "Read-only review link (#…&ro=1) — a READ-ONLY badge replaces the machine controls; the link "
        "opened straight into the Mold & Clamp context.")

    # ---------- 8. Responsive / visual design ----------
    H(doc, "8. Visual design", 1)
    P(doc, "A dark, single-theme industrial palette. The five subsystem hues are the only saturated "
           "colours in the chrome, so a lit subsystem reads instantly; a continuous green→amber→red "
           "scale encodes live health on the model and in the gauges. Numbers are set in a monospace "
           "with tabular figures so streaming values don't jitter. The layout is desktop-optimised and "
           "keeps the 3-D viewer the dominant column from 1280 px up.")
    palette_table(doc, [
        ("background", "#0a0e16", "app background"),
        ("surface", "#121826", "cards / panels"),
        ("elevated", "#1b2433", "header, dialogs, hover"),
        ("border", "#263043", "hairline borders"),
        ("text primary", "#e8edf4", "headings, key numbers"),
        ("accent", "#38bdf8", "interactive chrome, focus"),
        ("success", "#34d399", "healthy / good / monitor"),
        ("warning", "#fbbf24", "watch / acceptable"),
        ("critical", "#f4554e", "critical / waste"),
        ("Hydraulic", "#3b82f6", "subsystem identity"),
        ("Screw&CheckRing", "#f97316", "subsystem identity"),
        ("Drive", "#22c55e", "subsystem identity"),
        ("Heaters", "#eab308", "subsystem identity"),
        ("Mold", "#a855f7", "subsystem identity"),
    ])
    img(doc, "21_responsive_1280.png", "The layout at 1280 px — the viewer stays the dominant column.")

    # ---------- 9. Deployment ----------
    H(doc, "9. Deterministic deployment (clone → run → share)", 1)
    P(doc, "The platform is built to come up identically on any machine — laptop or DGX — with no "
           "manual recreation and no repeated mapping or training.")
    P(doc, "Committed to the repository (survives a clone):", bold=True)
    bullet(doc, "All code, the 3-D model (model.dae) and its textures.")
    bullet(doc, "The component map — public/map/component-map.detailed.json (source of truth).")
    bullet(doc, "Trained ML models (artifacts/models/*.pkl) so no retraining is needed.")
    bullet(doc, "Default platform config (config/platform.json) and the deterministic simulator seed.")
    P(doc, "Runtime persistence is to files, written atomically (temp + rename) and validated on load. "
           "The server serves the platform build by default; setting WEB_DIR=web/dist rolls back to the "
           "legacy dashboard instantly.")
    P(doc, "On the DGX:", bold=True)
    code(doc,
         "git clone <repo> && cd injection-molding\n"
         "./demo.sh            # build + serve at http://localhost:8000\n"
         "./demo.sh --tunnel   # also print a temporary public https link (Cloudflare)\n"
         "#   relaunch any time with the same command; first run does one-time setup only")
    P(doc, "demo.sh is idempotent: it creates the venv, trains models only if missing, builds the "
           "platform frontend only if missing, then serves; --tunnel adds a shareable public URL for a "
           "remote investor demo.")

    # ---------- 10. Glossary ----------
    H(doc, "10. Glossary — plain English", 1)
    glossary_entry(doc, "Digital twin",
        "A live software model of the machine driven by the same data stream — here an interactive 3-D "
        "model tinted by each subsystem's real-time health.",
        "A virtual copy of the press that lights up where it hurts.",
        "A dashboard stand-in of your machine that shows condition on the actual geometry.")
    glossary_entry(doc, "Subsystem",
        "One of five monitored asset groups (Hydraulic, Screw & Check Ring, Drive, Heaters, Mold & "
        "Clamp), each mapped to a region of the model and a set of sensors.",
        "A major part of the machine that wears and gets replaced as a unit.",
        "The pump pack, the screw, the motor, the heaters, the mold — the things you actually service.")
    glossary_entry(doc, "RUL (Remaining Useful Life)",
        "The predicted number of production cycles a subsystem will survive before its health crosses "
        "the failure threshold.",
        "How much life the part has left, counted in shots.",
        "Like an oil-life gauge — but counting molded parts instead of miles.")
    glossary_entry(doc, "p10 / p50 / p90",
        "Three points on the predicted RUL distribution — a pessimistic, most-likely and optimistic "
        "estimate; the band on the forecast chart.",
        "A best-case, likely, and worst-case estimate of remaining life.",
        "'Most likely 2,000 more shots, but could be fewer or more' — shown as the shaded band.")
    glossary_entry(doc, "Urgency band",
        "A bucket (critical / imminent / schedule / monitor) derived from days-until-replacement.",
        "A traffic-light label for how soon to act on a part.",
        "Red = order it now; green = check again next month.")
    glossary_entry(doc, "Machine state",
        "A ground-truth status (Running / Warning / Critical / Failed) from the worst subsystem's true "
        "health — distinct from the ML-predicted urgency.",
        "The overall traffic-light for the whole press, right now.",
        "Green = making good parts; red FAILED = stopped, reset to continue.")
    glossary_entry(doc, "Component map",
        "The file that maps each mesh of the 3-D model to a subsystem, so live health can tint the "
        "right geometry.",
        "The link between the picture of the machine and the data about it.",
        "What makes the mold end light up purple and the pump pack light up blue.")
    glossary_entry(doc, "Fault injection",
        "A demo control that ramps extra wear into a subsystem to simulate a developing fault.",
        "A test button that fakes a part going bad.",
        "Like a trainer tripping a sensor to prove the alarm catches it — demonstration only.")

    # ---------- 11. Roadmap ----------
    H(doc, "11. Roadmap", 1)
    bullet(doc, "Connect a real machine through the PLC DataSource (the snapshot contract is unchanged) "
                "and validate predictions against run-to-failure data.")
    bullet(doc, "Expert-refined component map per machine variant, versioned in the repo.")
    bullet(doc, "Threshold overrides and per-machine configuration via the persisted config API.")
    bullet(doc, "Cursor-level multiplayer review (presence) on top of the shareable-link sessions.")
    bullet(doc, "Fleet view across many machines, backed by a historian/database for trends and audit.")
    bullet(doc, "Authentication and a read-only reviewer role for remote investor / customer demos.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"wrote {OUT}")


if __name__ == "__main__":
    main()
